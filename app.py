from flask import Flask, render_template, request, jsonify, redirect, url_for, send_from_directory
import os
import io
from pypdf import PdfReader
from docx import Document
from pdf2image import convert_from_path
from PIL import Image
import pythoncom
import tempfile
from werkzeug.utils import secure_filename
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['ALLOWED_EXTENSIONS'] = {'pdf', 'docx', 'txt'}
app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024  # 10MB max file size
app.config['IMAGE_FOLDER'] = 'static/images'

# Buat folder yang diperlukan
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['IMAGE_FOLDER'], exist_ok=True)

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

def extract_text_from_pdf(filepath):
    """Ekstrak teks dari file PDF"""
    text = ""
    try:
        with open(filepath, 'rb') as file:
            reader = PdfReader(file)
            for page in reader.pages:
                text += page.extract_text() + "\n"
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
    return text.strip()

def extract_images_from_pdf(filepath, output_folder, filename_prefix):
    """Ekstrak gambar dari file PDF"""
    try:
        images = convert_from_path(filepath)
        image_paths = []
        
        for i, image in enumerate(images):
            image_name = f"{filename_prefix}_page_{i+1}.jpg"
            image_path = os.path.join(output_folder, image_name)
            image.save(image_path, 'JPEG')
            image_paths.append(f"/{app.config['IMAGE_FOLDER']}/{image_name}")
            
        return image_paths
    except Exception as e:
        print(f"Error extracting images from PDF: {e}")
        return []

def extract_text_from_docx(filepath):
    """Ekstrak teks dari file DOCX"""
    try:
        doc = Document(filepath)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text])
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
        return ""

def extract_images_from_docx(filepath, output_folder, filename_prefix):
    """Ekstrak gambar dari file DOCX"""
    try:
        doc = Document(filepath)
        image_paths = []
        
        for i, rel in enumerate(doc.part.rels.values()):
            if "image" in rel.reltype:
                image_data = rel.target_part.blob
                image_name = f"{filename_prefix}_image_{i+1}.png"
                image_path = os.path.join(output_folder, image_name)
                
                with open(image_path, 'wb') as f:
                    f.write(image_data)
                
                # Konversi ke format yang lebih ringan
                img = Image.open(image_path)
                if img.mode in ('RGBA', 'P'):
                    img = img.convert('RGB')
                img.save(image_path, 'JPEG', quality=85)
                
                image_paths.append(f"/{app.config['IMAGE_FOLDER']}/{image_name}")
                
        return image_paths
    except Exception as e:
        print(f"Error extracting images from DOCX: {e}")
        return []

def extract_text_from_txt(filepath):
    """Baca teks dari file TXT"""
    try:
        with open(filepath, 'r', encoding='utf-8') as file:
            return file.read()
    except Exception as e:
        print(f"Error reading text file: {e}")
        return ""

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/uploads/<filename>')
def uploaded_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

@app.route('/static/images/<filename>')
def serve_image(filename):
    return send_from_directory(app.config['IMAGE_FOLDER'], filename)

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'Tidak ada file yang dipilih'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'Nama file kosong'}), 400
    
    if file and allowed_file(file.filename):
        # Inisialisasi pythoncom untuk ekstraksi gambar dari DOCX
        pythoncom.CoInitialize()
        
        try:
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            
            # Ekstrak teks
            if filename.lower().endswith('.pdf'):
                text = extract_text_from_pdf(filepath)
                images = extract_images_from_pdf(filepath, app.config['IMAGE_FOLDER'], 
                                               os.path.splitext(filename)[0])
            elif filename.lower().endswith('.docx'):
                text = extract_text_from_docx(filepath)
                images = extract_images_from_docx(filepath, app.config['IMAGE_FOLDER'],
                                                os.path.splitext(filename)[0])
            else:  # txt
                text = extract_text_from_txt(filepath)
                images = []
            
            # Hapus file asli setelah diproses
            # os.remove(filepath)
            
            return jsonify({
                'message': 'File berhasil diproses',
                'filename': filename,
                'text': text[:1000] + '...' if len(text) > 1000 else text,
                'text_length': len(text),
                'images': images
            })
            
        except Exception as e:
            return jsonify({'error': f'Terjadi kesalahan: {str(e)}'}), 500
            
        finally:
            # Pastikan untuk membersihkan COM objects
            pythoncom.CoUninitialize()
    
    return jsonify({'error': 'Format file tidak didukung'}), 400

if __name__ == '__main__':
    app.run(debug=True)